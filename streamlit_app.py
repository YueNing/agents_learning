import streamlit as st
import requests
import time
import json
import io
from datetime import datetime
import logging
from pathlib import Path
from typing import Dict, Optional
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from mutagen import File as MutagenFile
import numpy as np
import subprocess
import os

# Setup logging
log_dir = Path("logs")
log_dir.mkdir(exist_ok=True)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/frontend.log'),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger("streamlit")

# Initialize session state
if 'transcription' not in st.session_state:
    st.session_state.transcription = None
if 'summary' not in st.session_state:
    st.session_state.summary = None
if 'processing_time' not in st.session_state:
    st.session_state.processing_time = None
if 'file_info' not in st.session_state:
    st.session_state.file_info = None
if 'dark_mode' not in st.session_state:
    st.session_state.dark_mode = False
if 'processing_stage' not in st.session_state:
    st.session_state.processing_stage = None

# Custom CSS for modern design
def load_css():
    """Load custom CSS for modern design with theme support"""
    # Base styles
    base_css = """
    <style>
    /* Import Google Fonts */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    /* Global styling */
    .main > div {
        font-family: 'Inter', sans-serif;
    }
    
    /* Modern cards */
    .modern-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1.5rem;
        border-radius: 16px;
        box-shadow: 0 10px 25px rgba(0,0,0,0.1);
        margin: 1rem 0;
        color: white;
        border: 1px solid rgba(255,255,255,0.1);
    }
    
    .stats-card {
        background: rgba(255,255,255,0.95);
        padding: 1.5rem;
        border-radius: 12px;
        box-shadow: 0 4px 15px rgba(0,0,0,0.08);
        margin: 0.5rem 0;
        border-left: 4px solid #667eea;
    }
    
    .progress-card {
        background: linear-gradient(135deg, #ff9a9e 0%, #fecfef 50%, #fecfef 100%);
        padding: 2rem;
        border-radius: 16px;
        box-shadow: 0 10px 25px rgba(0,0,0,0.1);
        margin: 1rem 0;
        text-align: center;
    }
    
    /* Animated elements */
    .pulse {
        animation: pulse 2s infinite;
    }
    
    @keyframes pulse {
        0% { opacity: 1; }
        50% { opacity: 0.5; }
        100% { opacity: 1; }
    }
    
    .fade-in {
        animation: fadeIn 0.5s ease-in;
    }
    
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(20px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    /* Modern buttons */
    .stButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 12px;
        padding: 0.5rem 2rem;
        font-weight: 600;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(102, 126, 234, 0.6);
    }
    
    /* Status indicators */
    .status-success {
        background: linear-gradient(135deg, #56ab2f 0%, #a8e6cf 100%);
        padding: 1rem;
        border-radius: 12px;
        color: white;
        text-align: center;
        font-weight: 600;
    }
    
    .status-processing {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        padding: 1rem;
        border-radius: 12px;
        color: white;
        text-align: center;
        font-weight: 600;
    }
    
    .status-error {
        background: linear-gradient(135deg, #fc466b 0%, #3f5efb 100%);
        padding: 1rem;
        border-radius: 12px;
        color: white;
        text-align: center;
        font-weight: 600;
    }
    
    /* Hide default streamlit elements */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Custom scrollbar */
    ::-webkit-scrollbar {
        width: 8px;
    }
    
    ::-webkit-scrollbar-track {
        background: #f1f1f1;
        border-radius: 4px;
    }
    
    ::-webkit-scrollbar-thumb {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 4px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: linear-gradient(135deg, #764ba2 0%, #667eea 100%);
    }
    """
    
    # Dark mode specific styles
    dark_mode_css = """
    /* Dark mode overrides */
    .stApp {
        background-color: #1e1e1e !important;
        color: #ffffff !important;
    }
    
    .main .block-container {
        background-color: #1e1e1e !important;
        color: #ffffff !important;
    }
    
    .stats-card {
        background: rgba(45,45,45,0.95) !important;
        color: #ffffff !important;
        border-left: 4px solid #667eea;
    }
    
    .stTextArea textarea {
        background-color: #2d2d2d !important;
        color: #ffffff !important;
        border: 1px solid #444 !important;
    }
    
    .stSelectbox select {
        background-color: #2d2d2d !important;
        color: #ffffff !important;
        border: 1px solid #444 !important;
    }
    
    .stMetric {
        background-color: #2d2d2d !important;
        color: #ffffff !important;
    }
    
    .stTabs [data-baseweb="tab-list"] {
        background-color: #2d2d2d !important;
    }
    
    .stTabs [data-baseweb="tab"] {
        background-color: #2d2d2d !important;
        color: #ffffff !important;
    }
    """
    
    # Combine CSS based on theme
    if st.session_state.dark_mode:
        combined_css = base_css + dark_mode_css + "</style>"
    else:
        combined_css = base_css + "</style>"
    
    st.markdown(combined_css, unsafe_allow_html=True)

def get_file_info(uploaded_file) -> Dict:
    """Extract comprehensive file information"""
    try:
        # Save file temporarily to analyze with mutagen
        temp_path = f"/tmp/{uploaded_file.name}"
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getvalue())
        
        # Get audio metadata using mutagen
        audio_file = MutagenFile(temp_path)
        duration = 0
        bitrate = 0
        
        if audio_file is not None and hasattr(audio_file, 'info'):
            # Handle different audio formats
            info = audio_file.info
            
            # Get duration
            if hasattr(info, 'length'):
                duration = info.length
            elif hasattr(info, 'duration'):
                duration = info.duration
            
            # Get bitrate
            if hasattr(info, 'bitrate'):
                bitrate = info.bitrate
            elif hasattr(info, 'total_bitrate'):
                bitrate = info.total_bitrate
            
            # For MP4/M4A files, try alternative methods
            if duration == 0 and uploaded_file.name.lower().endswith(('.m4a', '.mp4', '.aac')):
                try:
                    from mutagen.mp4 import MP4
                    mp4_file = MP4(temp_path)
                    if mp4_file.info:
                        duration = mp4_file.info.length
                        bitrate = mp4_file.info.bitrate
                except:
                    pass
        
        # If still no duration, try ffprobe as fallback
        if duration == 0:
            try:
                # Try using ffprobe to get accurate metadata
                result = subprocess.run([
                    'ffprobe', '-v', 'quiet', '-show_entries', 
                    'format=duration,bit_rate', '-of', 'csv=p=0', temp_path
                ], capture_output=True, text=True, timeout=30)
                
                if result.returncode == 0:
                    lines = result.stdout.strip().split('\n')
                    if len(lines) > 0:
                        parts = lines[0].split(',')
                        if len(parts) >= 1 and parts[0]:
                            duration = float(parts[0])
                        if len(parts) >= 2 and parts[1]:
                            bitrate = float(parts[1])
                        logger.info(f"Successfully extracted metadata using ffprobe for {uploaded_file.name}")
            except Exception as ffprobe_error:
                logger.warning(f"ffprobe failed: {ffprobe_error}")
        
        # If still no duration, estimate based on file size and typical bitrates
        if duration == 0:
            file_size_mb = len(uploaded_file.getvalue()) / (1024 * 1024)
            # Estimate duration based on typical audio bitrates (128-320 kbps average)
            estimated_bitrate = 192  # kbps average
            duration = (file_size_mb * 8 * 1024) / estimated_bitrate  # Convert MB to seconds
            bitrate = estimated_bitrate * 1000  # Convert to bps
            logger.warning(f"Could not extract metadata for {uploaded_file.name}, using size-based estimates")
        
        file_size = len(uploaded_file.getvalue())
        
        return {
            "name": uploaded_file.name,
            "size_bytes": file_size,
            "size_mb": file_size / (1024 * 1024),
            "duration_seconds": duration,
            "duration_minutes": duration / 60 if duration > 0 else 0,
            "bitrate_kbps": bitrate // 1000 if bitrate > 0 else 0,
            "format": uploaded_file.name.split('.')[-1].upper(),
            "estimated_words": int(duration * 2.5) if duration > 0 else 0,  # ~150 words per minute / 60 seconds
            "upload_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
    except Exception as e:
        logger.error(f"Error getting file info: {e}")
        
        # Fallback with size-based estimation
        file_size = len(uploaded_file.getvalue())
        file_size_mb = file_size / (1024 * 1024)
        
        # Rough estimation for audio files based on typical compression
        estimated_duration = (file_size_mb * 8 * 1024) / 192  # 192 kbps average
        
        return {
            "name": uploaded_file.name,
            "size_bytes": file_size,
            "size_mb": file_size_mb,
            "duration_seconds": estimated_duration,
            "duration_minutes": estimated_duration / 60,
            "bitrate_kbps": 192,  # Estimated average
            "format": uploaded_file.name.split('.')[-1].upper(),
            "estimated_words": int(estimated_duration * 2.5),
            "upload_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }

def estimate_processing_time(file_info: Dict) -> Dict:
    """Estimate processing times based on file characteristics"""
    duration = file_info.get("duration_minutes", 0)
    size_mb = file_info.get("size_mb", 0)
    
    # Rough estimates (adjust based on your hardware)
    transcription_time = max(duration * 0.3, size_mb * 0.1)  # 30% of audio duration or 0.1 min per MB
    summarization_time = max(duration * 0.05, 0.5)  # 5% of audio duration or minimum 30 seconds
    
    return {
        "transcription_estimate": transcription_time,
        "summarization_estimate": summarization_time,
        "total_estimate": transcription_time + summarization_time
    }

def display_file_info_card(file_info: Dict):
    """Display file information in a modern card"""
    st.markdown("""
    <div class="stats-card fade-in">
        <h3 style="margin-top: 0; color: #667eea;">üìÅ File Information</h3>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("üìÑ Format", file_info.get("format", "Unknown"))
        st.metric("‚è±Ô∏è Duration", f"{file_info.get('duration_minutes', 0):.1f} min")
    
    with col2:
        st.metric("üíæ Size", f"{file_info.get('size_mb', 0):.1f} MB")
        st.metric("üéµ Bitrate", f"{file_info.get('bitrate_kbps', 0)} kbps")
    
    with col3:
        st.metric("üìù Est. Words", f"{file_info.get('estimated_words', 0):,}")
        st.metric("üìÖ Uploaded", file_info.get("upload_time", "Unknown"))
    
    with col4:
        estimates = estimate_processing_time(file_info)
        st.metric("‚è≥ Est. Processing", f"{estimates.get('total_estimate', 0):.1f} min")
        st.metric("ü§ñ Model", "Whisper Large")

def display_processing_stats(processing_time: Dict, file_info: Dict):
    """Display processing statistics with charts"""
    st.markdown("""
    <div class="stats-card fade-in">
        <h3 style="margin-top: 0; color: #667eea;">üìä Processing Statistics</h3>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Processing time breakdown chart
        times = {
            'Transcription': processing_time.get('transcription', 0),
            'Summarization': processing_time.get('summarization', 0)
        }
        
        fig = go.Figure(data=[
            go.Bar(
                x=list(times.keys()),
                y=list(times.values()),
                marker_color=['#667eea', '#764ba2'],
                text=[f"{v:.1f}s" for v in times.values()],
                textposition='auto',
            )
        ])
        
        fig.update_layout(
            title="Processing Time Breakdown",
            xaxis_title="Stage",
            yaxis_title="Time (seconds)",
            template="plotly_white",
            height=300
        )
        
        st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        # Performance metrics with proper error handling
        duration_min = file_info.get('duration_minutes', 0)
        total_time = processing_time.get('total', 0)
        estimated_words = file_info.get('estimated_words', 0)
        
        # Calculate metrics with safe division
        if total_time > 0 and duration_min > 0:
            processing_speed = duration_min / (total_time / 60)
            words_per_second = estimated_words / total_time
            efficiency = processing_speed * 100
        else:
            processing_speed = 0
            words_per_second = 0
            efficiency = 0
        
        metrics_data = {
            'Processing Speed': f"{processing_speed:.1f}x real-time" if processing_speed > 0 else "N/A",
            'Words per Second': f"{words_per_second:.1f}" if words_per_second > 0 else "N/A",
            'Efficiency': f"{efficiency:.0f}%" if efficiency > 0 else "N/A"
        }
        
        st.markdown("### ‚ö° Performance Metrics")
        for metric, value in metrics_data.items():
            st.metric(metric, value)

def create_progress_indicator(stage: str, progress: int):
    """Create animated progress indicator"""
    stage_info = {
        "uploading": {"emoji": "üì§", "text": "Uploading audio file", "color": "#ff9a9e"},
        "transcribing": {"emoji": "üé§", "text": "Converting speech to text", "color": "#667eea"},
        "summarizing": {"emoji": "üß†", "text": "Generating intelligent summary", "color": "#764ba2"},
        "completed": {"emoji": "‚úÖ", "text": "Processing completed successfully", "color": "#56ab2f"}
    }
    
    info = stage_info.get(stage, {"emoji": "‚è≥", "text": "Processing", "color": "#667eea"})
    
    st.markdown(f"""
    <div class="progress-card pulse">
        <h2 style="margin: 0; font-size: 3rem;">{info['emoji']}</h2>
        <h3 style="margin: 0.5rem 0; color: #333;">{info['text']}</h3>
        <p style="margin: 0; color: #666;">Progress: {progress}%</p>
    </div>
    """, unsafe_allow_html=True)

def export_full_report_to_pdf(transcription: str, summary: Dict, file_info: Dict) -> bytes:
    """Export transcription and summary to PDF"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title = Paragraph("Audio Transcription & Summary Report", styles['Title'])
    story.append(title)
    story.append(Spacer(1, 12))
    
    # File info
    file_info_text = f"""
    <b>File:</b> {file_info.get('name', 'Unknown')}<br/>
    <b>Duration:</b> {file_info.get('duration_minutes', 0):.1f} minutes<br/>
    <b>Size:</b> {file_info.get('size_mb', 0):.1f} MB<br/>
    <b>Processed:</b> {file_info.get('upload_time', 'Unknown')}
    """
    story.append(Paragraph(file_info_text, styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Transcription
    story.append(Paragraph("Transcription", styles['Heading1']))
    story.append(Paragraph(transcription, styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Summary
    story.append(Paragraph("Summary", styles['Heading1']))
    
    # Check if full_text exists and use it as the primary content
    if summary.get('full_text'):
        # Split the full_text into paragraphs and format them properly
        full_text_content = summary['full_text']
        # Replace line breaks and format for PDF
        full_text_content = full_text_content.replace('\n', '<br/>')
        story.append(Paragraph(full_text_content, styles['Normal']))
        story.append(Spacer(1, 20))
    else:
        # Fallback to structured sections if full_text is not available
        if summary.get('overview'):
            story.append(Paragraph(f"<b>Overview:</b> {summary['overview']}", styles['Normal']))
            story.append(Spacer(1, 12))
    
        for section_name, section_data in summary.items():
            if section_name not in ['overview', 'full_text'] and section_data:
                story.append(Paragraph(f"<b>{section_name.replace('_', ' ').title()}:</b>", styles['Heading2']))
                if isinstance(section_data, list):
                    for item in section_data:
                        story.append(Paragraph(f"‚Ä¢ {item}", styles['Normal']))
                else:
                    story.append(Paragraph(str(section_data), styles['Normal']))
                story.append(Spacer(1, 12))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def export_summary_to_pdf(summary: Dict, file_info: Dict) -> bytes:
    """Export only summary to PDF"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title = Paragraph("Audio Summary Report", styles['Title'])
    story.append(title)
    story.append(Spacer(1, 12))
    
    # File info
    file_info_text = f"""
    <b>File:</b> {file_info.get('name', 'Unknown')}<br/>
    <b>Duration:</b> {file_info.get('duration_minutes', 0):.1f} minutes<br/>
    <b>Size:</b> {file_info.get('size_mb', 0):.1f} MB<br/>
    <b>Processed:</b> {file_info.get('upload_time', 'Unknown')}
    """
    story.append(Paragraph(file_info_text, styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Summary Content - Include full_text if available
    story.append(Paragraph("Summary", styles['Heading1']))
    
    # Check if full_text exists and use it as the primary content
    if summary.get('full_text'):
        # Split the full_text into paragraphs and format them properly
        full_text_content = summary['full_text']
        # Replace line breaks and format for PDF
        full_text_content = full_text_content.replace('\n', '<br/>')
        story.append(Paragraph(full_text_content, styles['Normal']))
        story.append(Spacer(1, 20))
    else:
        # Fallback to structured sections if full_text is not available
        if summary.get('overview'):
            story.append(Paragraph(f"<b>Overview:</b> {summary['overview']}", styles['Normal']))
            story.append(Spacer(1, 12))
        
        for section_name, section_data in summary.items():
            if section_name not in ['overview', 'full_text'] and section_data:
                story.append(Paragraph(f"<b>{section_name.replace('_', ' ').title()}:</b>", styles['Heading2']))
                if isinstance(section_data, list):
                    for item in section_data:
                        story.append(Paragraph(f"‚Ä¢ {item}", styles['Normal']))
                else:
                    story.append(Paragraph(str(section_data), styles['Normal']))
                story.append(Spacer(1, 12))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def export_to_word(transcription: str, summary: Dict, file_info: Dict) -> bytes:
    """Export transcription and summary to Word document"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Audio Transcription & Summary Report', 0)
    
    # File info
    doc.add_heading('File Information', level=1)
    file_info_para = doc.add_paragraph()
    file_info_para.add_run(f"File: ").bold = True
    file_info_para.add_run(f"{file_info.get('name', 'Unknown')}\n")
    file_info_para.add_run(f"Duration: ").bold = True
    file_info_para.add_run(f"{file_info.get('duration_minutes', 0):.1f} minutes\n")
    file_info_para.add_run(f"Size: ").bold = True
    file_info_para.add_run(f"{file_info.get('size_mb', 0):.1f} MB\n")
    file_info_para.add_run(f"Processed: ").bold = True
    file_info_para.add_run(f"{file_info.get('upload_time', 'Unknown')}")
    
    # Transcription
    doc.add_heading('Transcription', level=1)
    doc.add_paragraph(transcription)
    
    # Summary
    doc.add_heading('Summary', level=1)
    if summary.get('overview'):
        overview_para = doc.add_paragraph()
        overview_para.add_run('Overview: ').bold = True
        overview_para.add_run(summary['overview'])
    
    for section_name, section_data in summary.items():
        if section_name not in ['overview', 'full_text'] and section_data:
            doc.add_heading(section_name.replace('_', ' ').title(), level=2)
            if isinstance(section_data, list):
                for item in section_data:
                    doc.add_paragraph(f"‚Ä¢ {item}")
            else:
                doc.add_paragraph(str(section_data))
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def display_summary(summary: Dict):
    """Display the structured summary with modern styling"""
    if not summary:
        st.warning("No summary available")
        return

    try:
        st.markdown("""
        <div class="stats-card fade-in">
            <h3 style="margin-top: 0; color: #667eea;">üìã Intelligent Summary</h3>
        </div>
        """, unsafe_allow_html=True)
        
        # Export buttons - Fixed to export only summary
        col1, col2, col3 = st.columns([1, 1, 2])
        with col1:
            if st.button("üìÑ Export Summary PDF", key="export_summary_pdf"):
                try:
                    pdf_data = export_summary_to_pdf(
                        summary,
                        st.session_state.file_info or {}
                    )
                    st.download_button(
                        "‚¨áÔ∏è Download Summary PDF",
                        data=pdf_data,
                        file_name=f"summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                        mime="application/pdf"
                    )
                except Exception as e:
                    st.error(f"PDF export failed: {e}")
        
        with col2:
            if st.button("üìù Export Word", key="export_word"):
                try:
                    word_data = export_to_word(
                        st.session_state.transcription or "",
                        summary,
                        st.session_state.file_info or {}
                    )
                    st.download_button(
                        "‚¨áÔ∏è Download Word",
                        data=word_data,
                        file_name=f"transcription_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Word export failed: {e}")

        # Overview Section
        if summary.get("overview"):
            st.markdown("### üéØ Overview")
            st.info(summary["overview"])

        # Create two columns for summary sections
        col1, col2 = st.columns(2)
        
        with col1:
            # Main Points Section
            if summary.get("main_points"):
                st.markdown("### üìå Main Points")
                for i, point in enumerate(summary["main_points"], 1):
                    st.markdown(f"**{i}.** {point}")

            # Action Items / Decisions Section
            if summary.get("action_items_decisions"):
                st.markdown("### ‚úÖ Action Items / Decisions")
                for i, item in enumerate(summary["action_items_decisions"], 1):
                    st.markdown(f"**{i}.** {item}")
        
        with col2:
            # Key Insights Section
            if summary.get("key_insights"):
                st.markdown("### üí° Key Insights")
                for i, insight in enumerate(summary["key_insights"], 1):
                    st.markdown(f"**{i}.** {insight}")

            # Open Questions / Next Steps Section
            if summary.get("open_questions_next_steps"):
                st.markdown("### ‚ùì Open Questions / Next Steps")
                for i, question in enumerate(summary["open_questions_next_steps"], 1):
                    st.markdown(f"**{i}.** {question}")

        # Conclusions Section
        if summary.get("conclusions"):
            st.markdown("### üéØ Conclusions")
            for i, conclusion in enumerate(summary["conclusions"], 1):
                st.markdown(f"**{i}.** {conclusion}")

        # Full Text at bottom
        if summary.get("full_text"):
            with st.expander("üìÑ Show Full Summary Text"):
                st.markdown(summary["full_text"])

    except Exception as e:
        st.error(f"Error displaying summary: {e}")
        logger.error(f"Error in display_summary: {e}")

def main():
    # Page configuration
    st.set_page_config(
        page_title="üéôÔ∏è Audio Transcription & AI Summary",
        page_icon="üéôÔ∏è",
        layout="wide",
        initial_sidebar_state="collapsed"
    )
    
    # Load custom CSS
    load_css()
    
    # Header with theme toggle
    col1, col2, col3 = st.columns([2, 1, 1])
    with col1:
        st.markdown("""
        <h1 style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                   -webkit-background-clip: text; -webkit-text-fill-color: transparent; 
                   font-size: 3rem; margin: 0;">
        üéôÔ∏è Audio Transcription & AI Summary
        </h1>
        <p style="color: #666; font-size: 1.2rem; margin-top: 0;">
        Transform your audio into intelligent insights with AI-powered transcription and summarization
        </p>
        """, unsafe_allow_html=True)
    
    with col3:
        # Theme toggle - Fixed functionality
        theme_button_text = "üåô Dark Mode" if not st.session_state.dark_mode else "‚òÄÔ∏è Light Mode"
        if st.button(theme_button_text, key="theme_toggle"):
            st.session_state.dark_mode = not st.session_state.dark_mode
            st.rerun()
    
    # Model info card
    st.markdown("""
    <div class="modern-card">
        <h3 style="margin-top: 0;">ü§ñ AI Models Information</h3>
        <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 1rem;">
            <div>
                <strong>üé§ Transcription:</strong> OpenAI Whisper Large<br/>
                <small>State-of-the-art speech recognition with 99% accuracy</small>
            </div>
            <div>
                <strong>üß† Summarization:</strong> Microsoft Phi-4<br/>
                <small>Advanced language model for intelligent content analysis</small>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)

    # File upload section
    uploaded_file = st.file_uploader(
        "üìÅ Choose an audio file", 
        type=["mp3", "wav", "m4a", "flac", "ogg"],
        help="Supported formats: MP3, WAV, M4A, FLAC, OGG"
    )

    if uploaded_file:
        # Get and store file information
        if st.session_state.file_info is None:
            st.session_state.file_info = get_file_info(uploaded_file)
        
        # Display file information
        display_file_info_card(st.session_state.file_info)
        
        # Processing button
        if st.button("üöÄ Start Processing", type="primary"):
            try:
                start_time = time.time()
                logger.info("Starting audio processing")
                
                # Initialize progress tracking
                progress_container = st.empty()
                status_container = st.empty()

                # Stage 1: Uploading
                with progress_container.container():
                    create_progress_indicator("uploading", 10)
                
                files = {"file": uploaded_file}
                
                # Stage 2: API Request
                with progress_container.container():
                    create_progress_indicator("transcribing", 30)
                
                response = requests.post(
                    "http://app:8000/transcribe",
                    files=files,
                    timeout=7200
                )
                
                if response.status_code == 200:
                    data = response.json()
                    
                    # Stage 3: Summarizing
                    with progress_container.container():
                        create_progress_indicator("summarizing", 80)
                    
                    # Store results
                    st.session_state.transcription = data.get("transcription", "")
                    st.session_state.summary = data.get("summary", {})
                    st.session_state.processing_time = data.get("processing_time", {})
                    
                    # Stage 4: Completed
                    with progress_container.container():
                        create_progress_indicator("completed", 100)
                    
                    time.sleep(1)  # Show completion briefly
                    progress_container.empty()
                    
                    st.markdown("""
                    <div class="status-success fade-in">
                        ‚úÖ Processing completed successfully!
                    </div>
                    """, unsafe_allow_html=True)
                    
                    st.rerun()
                    
                else:
                    progress_container.empty()
                    st.markdown(f"""
                    <div class="status-error">
                        ‚ùå API Error: {response.status_code}<br/>
                        {response.text}
                    </div>
                    """, unsafe_allow_html=True)
                    
            except Exception as e:
                progress_container.empty()
                st.markdown(f"""
                <div class="status-error">
                    ‚ùå Processing Error: {str(e)}
                </div>
                """, unsafe_allow_html=True)
                logger.error(f"Processing error: {e}")

    # Display results if available
    if (hasattr(st.session_state, 'transcription') and 
        st.session_state.transcription and 
        hasattr(st.session_state, 'summary') and 
        st.session_state.summary):
        
        st.markdown("<br><br>", unsafe_allow_html=True)
        
        # Processing statistics
        if st.session_state.processing_time and st.session_state.file_info:
            display_processing_stats(st.session_state.processing_time, st.session_state.file_info)
        
        # Results tabs
        tab1, tab2 = st.tabs(["üìù Transcription", "üìã Summary"])
        
        with tab1:
            st.markdown("""
            <div class="stats-card fade-in">
                <h3 style="margin-top: 0; color: #667eea;">üìù Full Transcription</h3>
            </div>
            """, unsafe_allow_html=True)
            
            # Transcription statistics
            word_count = len(st.session_state.transcription.split())
            char_count = len(st.session_state.transcription)
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("üìä Word Count", f"{word_count:,}")
            with col2:
                st.metric("üî§ Character Count", f"{char_count:,}")
            with col3:
                if st.session_state.file_info:
                    duration = st.session_state.file_info.get('duration_minutes', 1)
                    wpm = word_count / duration if duration > 0 else 0
                    st.metric("‚ö° Words per Minute", f"{wpm:.0f}")
            with col4:
                if st.button("üìÑ Export Full Report PDF", key="export_full_pdf"):
                    try:
                        pdf_data = export_full_report_to_pdf(
                            st.session_state.transcription,
                            st.session_state.summary or {},
                            st.session_state.file_info or {}
                        )
                        st.download_button(
                            "‚¨áÔ∏è Download Full Report PDF",
                            data=pdf_data,
                            file_name=f"full_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf"
                        )
                    except Exception as e:
                        st.error(f"PDF export failed: {e}")
            
            # Download transcription button
            if st.button("üì• Download Transcription TXT", key="download_transcription"):
                st.download_button(
                    "‚¨áÔ∏è Download as TXT",
                    data=st.session_state.transcription,
                    file_name=f"transcription_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain"
                )
            
            st.text_area(
                "Full Transcription", 
                st.session_state.transcription, 
                height=400,
                help="Click and drag to select text for copying"
            )
            
        with tab2:
            if st.session_state.summary:
                display_summary(st.session_state.summary)
            else:
                st.warning("No summary available")

if __name__ == "__main__":
    main()
